"""ゴールデン（examples/before.docx → after.docx）に対する回帰テスト。

注意: examples/after.docx は旧バージョンで生成されており、インデントの数値
（CHAR_TWIPS や Chars 属性の有無、outlineLvl の有無）が現行コードと異なる。
そのため数値インデントの一致は検証せず、
  (1) 段落数の一致
  (2) 各段落テキストの一致（全角変換＋見出し再付番が正しく効いていること）
  (3) 見出し段落に outlineLvl が付く構造不変条件
を固定する。インデントの数値仕様は test_indent.py 側でユニットに固定する。
"""

from docx import Document
from docx.oxml.ns import qn


def _convert(examples_dir, tmp_path):
    from court_format_converter import convert
    out = tmp_path / "out.docx"
    convert(str(examples_dir / "before.docx"), str(out))
    return Document(str(out))


def test_paragraph_count_matches_golden(examples_dir, tmp_path):
    got = _convert(examples_dir, tmp_path)
    expected = Document(str(examples_dir / "after.docx"))
    assert len(got.paragraphs) == len(expected.paragraphs)


def test_paragraph_text_matches_golden(examples_dir, tmp_path):
    # テキスト（全角変換・見出し再付番後）はゴールデンと完全一致するはず
    got = _convert(examples_dir, tmp_path)
    expected = Document(str(examples_dir / "after.docx"))
    for i, (g, e) in enumerate(zip(got.paragraphs, expected.paragraphs)):
        assert g.text == e.text, f"段落{i}: {g.text!r} != {e.text!r}"


def test_headings_have_outline_level(examples_dir, tmp_path):
    # 「第１」「第２」など再付番された見出しに outlineLvl が付くこと
    got = _convert(examples_dir, tmp_path)
    heading_count = 0
    for p in got.paragraphs:
        pPr = p._element.find(qn("w:pPr"))
        if pPr is None:
            continue
        if pPr.find(qn("w:outlineLvl")) is not None:
            heading_count += 1
    # before.docx には 第１/１/２/(1)/(2)/第２ の6見出しがある
    assert heading_count >= 6


def test_no_halfwidth_kana_in_output(examples_dir, tmp_path):
    # 変換後の本文に半角カナ（U+FF61–U+FF9F）が残っていないこと
    got = _convert(examples_dir, tmp_path)
    for p in got.paragraphs:
        assert not any(0xFF61 <= ord(ch) <= 0xFF9F for ch in p.text), p.text
