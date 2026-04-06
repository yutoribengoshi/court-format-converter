#!/usr/bin/env python3
"""
裁判所書式変換ツール (court_format_converter.py)
===================================================
スタイル未登録のdocxファイルを裁判所書式（岡口マクロ準拠）に変換する。

使い方:
    python3 court_format_converter.py input.docx [output.docx]
    出力ファイル未指定時は「<元ファイル名>_裁判所書式.docx」に保存。

変換内容:
    - ページ設定: A4、余白(上35/下25/左30/右20mm)、26行×37文字グリッド
    - フォント: ＭＳ 明朝 / Times New Roman 12pt に統一
    - 見出し: テキストパターンから自動判定しインデント適用
    - テーブル: フォント統一＋レイアウト調整
    - フッター: ページ番号（中央）
    - 半角→全角変換: 数字・英字・カタカナ・括弧・記号を全角に統一
      （テーブル内・GPS座標等のデータ部分も含む）
"""

import sys
import re
import os
import datetime
from docx import Document
from docx.shared import Pt, Mm, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ============================================================
# 定数
# ============================================================

# 1全角文字 = 245 twips (12pt MS明朝基準)
CHAR_TWIPS = 245

# 見出しレベルごとの設定: (left_chars, 番号説明)
# 岡口マクロVBAソースの実際の値:
#   第１ = LeftIndent:24pt(2字), FirstLineIndent:-24pt(-2字) → 1行目0字目、2行目2字目
#   １   = LeftIndent:24pt(2字), FirstLineIndent:-12pt(-1字) → 1行目1字目、2行目2字目
#   (1)  = LeftIndent:36pt(3字), FirstLineIndent:-12pt(-1字) → 1行目2字目、2行目3字目
#   ア   = LeftIndent:48pt(4字), FirstLineIndent:-12pt(-1字) → 1行目3字目、2行目4字目
#   (ア) = LeftIndent:60pt(5字), FirstLineIndent:-12pt(-1字) → 1行目4字目、2行目5字目
#   ａ   = LeftIndent:72pt(6字), FirstLineIndent:-12pt(-1字) → 1行目5字目、2行目6字目
#   (a)  = LeftIndent:84pt(7字), FirstLineIndent:-12pt(-1字) → 1行目6字目、2行目7字目
    # 見出しの開始位置（leftの値）。ぶら下げは使わない。
    # 第１の「１」の真下に次の階層の「１」が来る。
HEADING_LEVELS = {
    1: (0, "第１"),    # 第=0字目、１=1字目
    2: (1, "１"),      # １=1字目（第１の１の真下）
    3: (2, "(1)"),     # (=2字目
    4: (3, "ア"),      # ア=3字目
    5: (4, "(ｱ)"),     # (=4字目
    6: (5, "ａ"),      # ａ=5字目
    7: (6, "(a)"),     # (=6字目
}

# 岡口マクロ準拠: 本文は見出しと同じLeftIndent + FirstLineIndent=12pt(1字)
BODY_INDENT = {
    # 本文は見出しの1字右 + 段落冒頭1字下げ（firstLine=1）。
    # 1行目 = left + 1字、2行目以降 = left。
    0: (1, 1),   # 見出しなし → 左1字 + 首行1字 → 1行目2字目、2行目1字目
    1: (1, 1),   # 第１直下 → 左1字 + 首行1字 → 1行目2字目、2行目1字目
    2: (2, 1),   # １直下 → 左2字 + 首行1字 → 1行目3字目、2行目2字目
    3: (3, 1),   # (1)直下 → 左3字 + 首行1字 → 1行目4字目、2行目3字目
    4: (4, 1),   # ア直下 → 左4字 + 首行1字 → 1行目5字目、2行目4字目
    5: (5, 1),   # (ｱ)直下 → 左5字 + 首行1字 → 1行目6字目、2行目5字目
    6: (6, 1),   # ａ直下 → 左6字 + 首行1字 → 1行目7字目、2行目6字目
    7: (7, 1),   # (a)直下 → 左7字 + 首行1字 → 1行目8字目、2行目7字目
}


# ============================================================
# 半角→全角変換
# ============================================================

# 半角→全角 変換テーブル
_HANKAKU_TO_ZENKAKU = str.maketrans(
    # 数字
    '0123456789'
    # 英大文字
    'ABCDEFGHIJKLMNOPQRSTUVWXYZ'
    # 英小文字
    'abcdefghijklmnopqrstuvwxyz'
    # 括弧・記号
    '()[]{}!?.,;:/-+=%&#@*~'
    # スペース（半角→全角）
    # ※スペースは変換しない（インデント崩れ防止）
    ,
    # 数字（全角）
    '０１２３４５６７８９'
    # 英大文字（全角）
    'ＡＢＣＤＥＦＧＨＩＪＫＬＭＮＯＰＱＲＳＴＵＶＷＸＹＺ'
    # 英小文字（全角）
    'ａｂｃｄｅｆｇｈｉｊｋｌｍｎｏｐｑｒｓｔｕｖｗｘｙｚ'
    # 括弧・記号（全角）
    '（）［］｛｝！？．，；：／−＋＝％＆＃＠＊〜'
)

# 半角カタカナ→全角カタカナ
_HANKAKU_KANA_MAP = {
    'ｱ': 'ア', 'ｲ': 'イ', 'ｳ': 'ウ', 'ｴ': 'エ', 'ｵ': 'オ',
    'ｶ': 'カ', 'ｷ': 'キ', 'ｸ': 'ク', 'ｹ': 'ケ', 'ｺ': 'コ',
    'ｻ': 'サ', 'ｼ': 'シ', 'ｽ': 'ス', 'ｾ': 'セ', 'ｿ': 'ソ',
    'ﾀ': 'タ', 'ﾁ': 'チ', 'ﾂ': 'ツ', 'ﾃ': 'テ', 'ﾄ': 'ト',
    'ﾅ': 'ナ', 'ﾆ': 'ニ', 'ﾇ': 'ヌ', 'ﾈ': 'ネ', 'ﾉ': 'ノ',
    'ﾊ': 'ハ', 'ﾋ': 'ヒ', 'ﾌ': 'フ', 'ﾍ': 'ヘ', 'ﾎ': 'ホ',
    'ﾏ': 'マ', 'ﾐ': 'ミ', 'ﾑ': 'ム', 'ﾒ': 'メ', 'ﾓ': 'モ',
    'ﾔ': 'ヤ', 'ﾕ': 'ユ', 'ﾖ': 'ヨ',
    'ﾗ': 'ラ', 'ﾘ': 'リ', 'ﾙ': 'ル', 'ﾚ': 'レ', 'ﾛ': 'ロ',
    'ﾜ': 'ワ', 'ﾝ': 'ン',
    'ﾞ': '゛', 'ﾟ': '゜', 'ｰ': 'ー', '｡': '。', '｢': '「', '｣': '」', '､': '、',
}

# 半角カタカナ濁点・半濁点の合成
_DAKUTEN_MAP = {
    'カ゛': 'ガ', 'キ゛': 'ギ', 'ク゛': 'グ', 'ケ゛': 'ゲ', 'コ゛': 'ゴ',
    'サ゛': 'ザ', 'シ゛': 'ジ', 'ス゛': 'ズ', 'セ゛': 'ゼ', 'ソ゛': 'ゾ',
    'タ゛': 'ダ', 'チ゛': 'ヂ', 'ツ゛': 'ヅ', 'テ゛': 'デ', 'ト゛': 'ド',
    'ハ゛': 'バ', 'ヒ゛': 'ビ', 'フ゛': 'ブ', 'ヘ゛': 'ベ', 'ホ゛': 'ボ',
    'ウ゛': 'ヴ',
    'ハ゜': 'パ', 'ヒ゜': 'ピ', 'フ゜': 'プ', 'ヘ゜': 'ペ', 'ホ゜': 'ポ',
}


def to_zenkaku(text):
    """半角文字を全角に変換する。"""
    # 1. 半角カタカナを全角に（濁点合成前に個別変換）
    result = []
    for ch in text:
        result.append(_HANKAKU_KANA_MAP.get(ch, ch))
    text = ''.join(result)

    # 2. 濁点・半濁点の合成
    for src, dst in _DAKUTEN_MAP.items():
        text = text.replace(src, dst)

    # 3. 数字・英字・記号の半角→全角
    text = text.translate(_HANKAKU_TO_ZENKAKU)

    return text


def convert_run_to_zenkaku(run):
    """run内のテキストを全角に変換。"""
    if run.text:
        run.text = to_zenkaku(run.text)


# ============================================================
# 見出し判定
# ============================================================

# 見出しパターン（優先順）
HEADING_PATTERNS = [
    # Level 1: 第１、第２、第１０ 等（全角・半角数字両対応）
    (1, re.compile(r'^[\s　]*第[１２３４５６７８９０\d]+[\s　]')),
    # Level 3: (1)、(2) 等（括弧＋数字）※Level 2より先に判定
    # スペースなしで直接テキストが続くケースにも対応
    (3, re.compile(r'^[\s　]*[\(（][１２３４５６７８９０\d]+[\)）][\s　]?')),
    # Level 5: (ｱ)、(ｲ) 等（カタカナ括弧）
    (5, re.compile(r'^[\s　]*[\(（][ｱ-ﾝア-ン]+[\)）][\s　]?')),
    # Level 7: (a)、(b) 等（括弧＋小文字）
    (7, re.compile(r'^[\s　]*[\(（][a-zａ-ｚ]+[\)）][\s　]?')),
    # Level 2: １、２、１０ 等（単独の全角数字 — 「第」なし）
    (2, re.compile(r'^[\s　]*[１２３４５６７８９０\d]+[\s　]')),
    # Level 4: ア、イ、ウ 等（単独の全角カタカナ1文字）
    (4, re.compile(r'^[\s　]*[ア-ン][　\s]')),
    # Level 6: ａ、ｂ 等（単独の全角小文字1文字）
    (6, re.compile(r'^[\s　]*[ａ-ｚ][　\s]')),
]

# 見出しとして扱わない短すぎる段落は除外しない（タイトルも短い）
# ただし「以上」「記」などの定型句は除外
SKIP_PATTERNS = re.compile(r'^[\s　]*(以上|記|別紙|添付|目録)[\s　]*$')

# 冒頭部分（事件番号、当事者名、表題、日付、弁護士名等）を判定
HEADER_PATTERNS = [
    re.compile(r'(原告|被告|申立人|被申立人|相手方|抗告人|債権者|債務者)'),
    re.compile(r'(準備書面|訴状|答弁書|意見書|報告書|申立書|陳述書|上申書)'),
    re.compile(r'(令和|平成|昭和)[０-９\d]+年'),
    re.compile(r'(弁護士|弁護人|代理人)'),
    re.compile(r'(裁判所|御[　\s]*中|殿)'),
    re.compile(r'(号証|甲|乙|丙)第?[０-９\d]'),
    re.compile(r'^[\s　]*(第[０-９\d]+[　\s])'),  # 最初の「第１」見出しもヘッダー後
]


def detect_heading_level(text):
    """段落テキストから見出しレベルを判定。見出しでなければ None を返す。"""
    stripped = text.strip().replace('\u3000', '　')

    if not stripped:
        return None

    if SKIP_PATTERNS.match(stripped):
        return None

    for level, pattern in HEADING_PATTERNS:
        if pattern.match(stripped):
            return level

    return None


def is_header_section(text):
    """冒頭セクション（事件番号〜弁護士名）かどうか判定。"""
    for pat in HEADER_PATTERNS:
        if pat.search(text):
            return True
    return False


    # 丸数字→（全角数字）に変換するマップ
_MARU_TO_KAKKO = {
    '⑴': '（１）', '⑵': '（２）', '⑶': '（３）', '⑷': '（４）', '⑸': '（５）',
    '⑹': '（６）', '⑺': '（７）', '⑻': '（８）', '⑼': '（９）', '⑽': '（１０）',
    '⑾': '（１１）', '⑿': '（１２）', '⒀': '（１３）', '⒁': '（１４）', '⒂': '（１５）',
    '⒃': '（１６）', '⒄': '（１７）', '⒅': '（１８）', '⒆': '（１９）', '⒇': '（２０）',
    '①': '（１）', '②': '（２）', '③': '（３）', '④': '（４）', '⑤': '（５）',
    '⑥': '（６）', '⑦': '（７）', '⑧': '（８）', '⑨': '（９）', '⑩': '（１０）',
    '⑪': '（１１）', '⑫': '（１２）', '⑬': '（１３）', '⑭': '（１４）', '⑮': '（１５）',
    '⑯': '（１６）', '⑰': '（１７）', '⑱': '（１８）', '⑲': '（１９）', '⑳': '（２０）',
}


def normalize_heading_spacing(text):
    """見出し番号の後のスペースを全角スペース1個に正規化する。
    丸数字（⑴⑵①②）は（１）（２）に変換。
    「（１）被疑者」→「（１）　被疑者」
    「（１）　　被疑者」→「（１）　被疑者」
    「１被告」→「１　被告」
    """
    ZS = '\u3000'  # 全角スペース

    # まず丸数字を（全角数字）に変換
    for maru, kakko in _MARU_TO_KAKKO.items():
        if text.lstrip().startswith(maru):
            leading = text[:len(text) - len(text.lstrip())]
            rest = text.lstrip()[len(maru):]
            text = leading + kakko + rest
            break

    # 括弧付き番号: （１）、（ｱ）、（a）等
    m = re.match(r'^([\s\u3000]*[\(（][１-９０-９\d]+[\)）])[\s\u3000]*(.*)', text, re.DOTALL)
    if m:
        return m.group(1) + ZS + m.group(2)
    m = re.match(r'^([\s\u3000]*[\(（][ｱ-ﾝア-ン]+[\)）])[\s\u3000]*(.*)', text, re.DOTALL)
    if m:
        return m.group(1) + ZS + m.group(2)
    m = re.match(r'^([\s\u3000]*[\(（][a-zａ-ｚ]+[\)）])[\s\u3000]*(.*)', text, re.DOTALL)
    if m:
        return m.group(1) + ZS + m.group(2)
    # 「第１」
    m = re.match(r'^([\s\u3000]*第[１-９０-９\d]+)[\s\u3000]*(.*)', text, re.DOTALL)
    if m:
        return m.group(1) + ZS + m.group(2)
    # 単独数字: 「１」「２」等
    m = re.match(r'^([\s\u3000]*[１-９０-９\d]+)[\s\u3000]*(.*)', text, re.DOTALL)
    if m:
        return m.group(1) + ZS + m.group(2)
    # 単独カタカナ: 「ア」「イ」等
    m = re.match(r'^([\s\u3000]*[ア-ン])[\s\u3000]*(.*)', text, re.DOTALL)
    if m:
        return m.group(1) + ZS + m.group(2)
    # 単独英字: 「ａ」「ｂ」等
    m = re.match(r'^([\s\u3000]*[ａ-ｚ])[\s\u3000]*(.*)', text, re.DOTALL)
    if m:
        return m.group(1) + ZS + m.group(2)
    return text


# ============================================================
# Wordコメント機能
# ============================================================

_comment_id_counter = 0

def _next_comment_id():
    global _comment_id_counter
    _comment_id_counter += 1
    return _comment_id_counter


def _ensure_comments_part(doc):
    """docにコメントパート（w:comments）がなければ作成。"""
    comments_part = None
    for rel in doc.part.rels.values():
        if "comments" in rel.reltype:
            comments_part = rel.target_part
            break

    if comments_part is None:
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI

        comments_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:comments xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"'
            ' xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"'
            ' xmlns:o="urn:schemas-microsoft-com:office:office"'
            ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
            ' xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"'
            ' xmlns:v="urn:schemas-microsoft-com:vml"'
            ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
            ' xmlns:w10="urn:schemas-microsoft-com:office:word"'
            ' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml"'
            '/>'
        )
        comments_part = Part(
            partname=PackURI('/word/comments.xml'),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml',
            blob=comments_xml.encode('utf-8'),
            package=doc.part.package,
        )
        doc.part.relate_to(
            comments_part,
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments',
        )

    return comments_part


def add_comment_to_paragraph(doc, para, comment_text, author="court-format"):
    """段落にWordコメントを追加する。"""
    comment_id = _next_comment_id()
    now = datetime.datetime.now().strftime('%Y-%m-%dT%H:%M:%S')

    # コメント本体をcommentsパートに追加
    comments_part = _ensure_comments_part(doc)
    from lxml import etree
    comments_elem = etree.fromstring(comments_part.blob)

    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    comment_elem = etree.SubElement(comments_elem, qn('w:comment'))
    comment_elem.set(qn('w:id'), str(comment_id))
    comment_elem.set(qn('w:author'), author)
    comment_elem.set(qn('w:date'), now + 'Z')

    # コメント内の段落
    comment_para = etree.SubElement(comment_elem, qn('w:p'))
    comment_run = etree.SubElement(comment_para, qn('w:r'))
    comment_t = etree.SubElement(comment_run, qn('w:t'))
    comment_t.text = comment_text
    comment_t.set(qn('xml:space'), 'preserve')

    comments_part._blob = etree.tostring(comments_elem, xml_declaration=True,
                                          encoding='UTF-8', standalone=True)

    # 段落にコメント参照マーカーを追加
    comment_range_start = parse_xml(
        f'<w:commentRangeStart {nsdecls("w")} w:id="{comment_id}"/>'
    )
    comment_range_end = parse_xml(
        f'<w:commentRangeEnd {nsdecls("w")} w:id="{comment_id}"/>'
    )
    comment_ref_run = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:rPr><w:rStyle w:val="CommentReference"/></w:rPr>'
        f'  <w:commentReference w:id="{comment_id}"/>'
        f'</w:r>'
    )

    # 段落の先頭にrangeStart、末尾にrangeEnd + reference
    para._element.insert(0, comment_range_start)
    para._element.append(comment_range_end)
    para._element.append(comment_ref_run)

    return comment_id


# ============================================================
# Word変更履歴（Track Changes）
# ============================================================

_revision_id_counter = 100  # コメントIDと衝突しないように100から開始


def _next_revision_id():
    global _revision_id_counter
    _revision_id_counter += 1
    return _revision_id_counter


def replace_paragraph_with_track_changes(para, new_text, author="court-format"):
    """段落のテキストを変更履歴付きで書き換える。
    Wordで開くと削除線（元テキスト）＋赤字（修正テキスト）として表示される。"""
    old_text = para.text
    if old_text == new_text:
        return False

    now = datetime.datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ')
    del_id = _next_revision_id()
    ins_id = _next_revision_id()

    # 既存のrunを全て削除
    for run in para.runs:
        run._element.getparent().remove(run._element)

    # 元テキストをw:del（削除線）で追加
    del_xml = (
        f'<w:del {nsdecls("w")} w:id="{del_id}" w:author="{author}" w:date="{now}">'
        f'  <w:r>'
        f'    <w:rPr>'
        f'      <w:rFonts w:eastAsia="\uff2d\uff33 \u660e\u671d" w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
        f'      <w:sz w:val="24"/>'
        f'    </w:rPr>'
        f'    <w:delText xml:space="preserve">{_escape_xml(old_text)}</w:delText>'
        f'  </w:r>'
        f'</w:del>'
    )
    para._element.append(parse_xml(del_xml))

    # 修正テキストをw:ins（挿入マーク）で追加
    ins_xml = (
        f'<w:ins {nsdecls("w")} w:id="{ins_id}" w:author="{author}" w:date="{now}">'
        f'  <w:r>'
        f'    <w:rPr>'
        f'      <w:rFonts w:eastAsia="\uff2d\uff33 \u660e\u671d" w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
        f'      <w:sz w:val="24"/>'
        f'    </w:rPr>'
        f'    <w:t xml:space="preserve">{_escape_xml(new_text)}</w:t>'
        f'  </w:r>'
        f'</w:ins>'
    )
    para._element.append(parse_xml(ins_xml))

    return True


def _escape_xml(text):
    """XMLの特殊文字をエスケープ。"""
    return (text
            .replace('&', '&amp;')
            .replace('<', '&lt;')
            .replace('>', '&gt;')
            .replace('"', '&quot;')
            .replace("'", '&apos;'))


# ============================================================
# フォント設定
# ============================================================

def set_run_font(run, size=12):
    """runのフォントをＭＳ明朝/Times New Roman に設定。"""
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rpr.insert(0, rfonts)
    rfonts.set(qn('w:eastAsia'), 'ＭＳ 明朝')
    rfonts.set(qn('w:ascii'), 'Times New Roman')
    rfonts.set(qn('w:hAnsi'), 'Times New Roman')


def set_paragraph_font(para, size=12):
    """段落内の全runのフォントを設定。"""
    for run in para.runs:
        set_run_font(run, size)


def _is_on_off_true(value):
    """WordprocessingML の on/off 値を真偽値に変換。"""
    return value is None or value.lower() in ('1', 'true', 'on')


def _is_fully_bold(paragraph_element):
    """段落内のテキストrunがすべて太字なら True。"""
    saw_text = False

    for run in paragraph_element.findall(qn('w:r')):
        text_nodes = run.findall(qn('w:t'))
        if not any(node.text for node in text_nodes):
            continue

        saw_text = True
        rpr = run.find(qn('w:rPr'))
        if rpr is None:
            return False

        bold = rpr.find(qn('w:b'))
        if bold is None or not _is_on_off_true(bold.get(qn('w:val'))):
            return False

    return saw_text


def apply_paragraph_layout(para, fmt):
    """元段落の余白・行間・太字を復元する。"""
    if not fmt:
        return

    pf = para.paragraph_format

    if fmt.get('space_before') is not None:
        pf.space_before = Twips(int(fmt['space_before']))
    if fmt.get('space_after') is not None:
        pf.space_after = Twips(int(fmt['space_after']))
    if fmt.get('line') is not None:
        pf.line_spacing = Twips(int(fmt['line']))
    if fmt.get('line_rule') == 'exact':
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    if fmt.get('bold'):
        for run in para.runs:
            run.bold = True


# ============================================================
# インデント設定
# ============================================================

def set_indent(para, left_chars=0, first_line_chars=0, hanging_chars=0):
    """段落にインデントを設定（文字単位）。
    hanging_chars: ぶら下げインデント（1行目を左に出す）。firstLineと排他。"""
    pPr = para._element.get_or_add_pPr()

    # 既存のindを削除
    existing_ind = pPr.find(qn('w:ind'))
    if existing_ind is not None:
        pPr.remove(existing_ind)

    if left_chars == 0 and first_line_chars == 0 and hanging_chars == 0:
        return

    ind = parse_xml(f'<w:ind {nsdecls("w")}/>')

    # Chars属性とtwips直値の両方を設定。
    # Chars属性が主、twipsはフォールバック（岡口マクロと同じ方式）。
    if left_chars > 0:
        ind.set(qn('w:leftChars'), str(int(left_chars * 100)))
        ind.set(qn('w:left'), str(int(left_chars * CHAR_TWIPS)))
    if hanging_chars > 0:
        ind.set(qn('w:hangingChars'), str(int(hanging_chars * 100)))
        ind.set(qn('w:hanging'), str(int(hanging_chars * CHAR_TWIPS)))
    elif first_line_chars > 0:
        ind.set(qn('w:firstLineChars'), str(int(first_line_chars * 100)))
        ind.set(qn('w:firstLine'), str(int(first_line_chars * CHAR_TWIPS)))

    pPr.append(ind)




def _set_outline_level(para, level):
    """段落にアウトラインレベルを設定（目次生成用）。
    見た目は変えずに、Wordの目次挿入機能で認識される。"""
    pPr = para._element.get_or_add_pPr()
    existing = pPr.find(qn('w:outlineLvl'))
    if existing is not None:
        pPr.remove(existing)
    # outlineLvl: 0=Level1, 1=Level2, ... 8=本文
    olvl = parse_xml(f'<w:outlineLvl {nsdecls("w")} w:val="{level - 1}"/>')
    pPr.append(olvl)


    # 見出しレベルごとのぶら下げ幅（岡口マクロVBAソースの実際の値）
    # 全レベル共通: FirstLineIndent = -12pt = -1字
    # 第１だけ例外: FirstLineIndent = -24pt = -2字
    # 岡口マクロVBAソース: FirstLineIndent値
    #   第１ = -24pt(-2字), それ以外 = -12pt(-1字)
_HEADING_HANGING = {
    1: 2,    # 第１ → 2字ぶら下げ（left2-hang2=0字目）
    2: 1,    # １　 → 1字ぶら下げ（left2-hang1=1字目）
    3: 1,    # (1)  → 1字ぶら下げ（left3字-hang1字=2字目から）
    4: 1,    # ア　 → 1字ぶら下げ（left4字-hang1字=3字目から）
    5: 1,    # (ア) → 1字ぶら下げ（left5字-hang1字=4字目から）
    6: 1,    # ａ　 → 1字ぶら下げ（left6字-hang1字=5字目から）
    7: 1,    # (a)  → 1字ぶら下げ（left7字-hang1字=6字目から）
}


def set_heading_indent(para, level):
    """見出し段落のインデント＋アウトラインレベル設定。
    leftのみで位置指定。ぶら下げは使わない（Wordの文字幅計算でずれるため）。"""
    left_chars = HEADING_LEVELS[level][0]
    _set_outline_level(para, level)
    set_indent(para, left_chars=left_chars)


def set_body_indent(para, current_heading_level, has_komidashi=True):
    """本文段落のインデント設定（直前の見出しレベルに基づく）。
    has_komidashi=True: 直前の見出しが小タイトル付き → firstLine=1（段落冒頭字下げ）
    has_komidashi=False: 直前の見出しが本文兼用 → firstLine=0（字下げなし）"""
    if current_heading_level in BODY_INDENT:
        left, fl = BODY_INDENT[current_heading_level]
    else:
        left, fl = 0, 1
    if not has_komidashi:
        fl = 0
    set_indent(para, left_chars=left, first_line_chars=fl)


# ============================================================
# ページ設定
# ============================================================

def setup_page(doc):
    """ページ設定を裁判所書式に変更。"""
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(35)
        section.bottom_margin = Mm(25)
        section.left_margin = Mm(30)
        section.right_margin = Mm(20)
        section.header_distance = Mm(0)
        section.footer_distance = Mm(15)

        # グリッド設定
        sectPr = section._sectPr
        docGrid = sectPr.find(qn('w:docGrid'))
        if docGrid is None:
            docGrid = parse_xml(
                f'<w:docGrid {nsdecls("w")} '
                f'w:type="linesAndChars" w:linePitch="516" w:charSpace="1057"/>'
            )
            sectPr.append(docGrid)
        else:
            docGrid.set(qn('w:type'), 'linesAndChars')
            docGrid.set(qn('w:linePitch'), '516')
            docGrid.set(qn('w:charSpace'), '1057')


# ============================================================
# ページ番号
# ============================================================

def add_page_number(doc):
    """フッター中央にページ番号を挿入。"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        # 既存のフッター内容をクリア
        for p in footer.paragraphs:
            p.clear()

        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # BEGIN field
        run1 = fp.add_run()
        set_run_font(run1, size=10)
        fld_begin = parse_xml(
            f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'
        )
        run1._element.append(fld_begin)

        # PAGE instruction
        run2 = fp.add_run()
        set_run_font(run2, size=10)
        instr = parse_xml(
            f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'
        )
        run2._element.append(instr)

        # 現在値を表示するため、separate + result を含む正規のフィールド構造にする。
        run3 = fp.add_run()
        set_run_font(run3, size=10)
        fld_sep = parse_xml(
            f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'
        )
        run3._element.append(fld_sep)

        run4 = fp.add_run('1')
        set_run_font(run4, size=10)

        # END field
        run5 = fp.add_run()
        set_run_font(run5, size=10)
        fld_end = parse_xml(
            f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'
        )
        run5._element.append(fld_end)


# ============================================================
# テーブル処理
# ============================================================

def format_tables(doc):
    """全テーブルの半角→全角変換＋フォントを統一。"""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        convert_run_to_zenkaku(run)
                        set_run_font(run, size=10)


# ============================================================
# デフォルトスタイル設定
# ============================================================

def setup_default_style(doc):
    """Normalスタイルのフォントを設定。"""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rpr.insert(0, rfonts)
    rfonts.set(qn('w:eastAsia'), 'ＭＳ 明朝')
    rfonts.set(qn('w:ascii'), 'Times New Roman')
    rfonts.set(qn('w:hAnsi'), 'Times New Roman')


# ============================================================
# LLM構造解析（--auto モード）
# ============================================================

AUTO_SYSTEM_PROMPT = """\
あなたは日本の裁判所提出書面の書式整形を行う専門家です。
与えられた文章の各段落を分析し、裁判所書式の見出し階層を判定してください。

## 見出し階層
- H1: 最上位の大見出し → 「第１　」「第２　」… を付与
- H2: 中見出し → 「１　」「２　」… を付与
- H3: 小見出し → 「(1)　」「(2)　」… を付与
- H4: 細目 → 「ア　」「イ　」… を付与
- H5: 細々目 → 「(ｱ)　」「(ｲ)　」… を付与
- H6: → 「ａ　」「ｂ　」… を付与
- H7: → 「(a)　」「(b)　」… を付与
- BODY: 本文（見出しではない）
- SKIP: 冒頭部分（事件番号、当事者名、表題、日付、署名等）・末尾の「以上」等 → そのまま

## ルール
1. 元の文章に見出し番号が既にあれば、それを正しい階層の全角番号に置き換える
2. 元の文章に見出し番号がなくても、意味的に見出しと判断できれば番号を振る
3. 各階層の番号は出現順に１から連番（H2はH1ごとにリセット、H3はH2ごとにリセット…）
4. 本文段落は BODY とする
5. 冒頭の事件番号・当事者表示・書面タイトル・日付・弁護士名等は SKIP とする
6. 末尾の「以上」も SKIP とする
7. 見出しテキストから元の番号を除去し、正しい番号を付与した完成テキストを返す

## 出力形式
各段落を1行ずつ、以下の形式で出力してください:
TYPE|テキスト

例:
SKIP|令和６年（ワ）第１２３号　損害賠償請求事件
SKIP|原告　山田太郎
SKIP|被告　株式会社○○
SKIP|準備書面（１）
SKIP|令和７年３月１日
SKIP|東京地方裁判所民事第○部　御中
SKIP|上記原告代理人
SKIP|弁護士　○○○○
H1|第１　被告の主張について
BODY|被告は令和５年１月に…
H2|１　契約の成立
BODY|原告と被告は…
H2|２　債務不履行
BODY|被告は上記契約に基づく…
H3|(1)　履行遅滞
BODY|被告は納期である…
H3|(2)　損害の発生
BODY|原告は上記履行遅滞により…
H1|第２　損害額の算定
BODY|原告の損害額は以下のとおりである。
SKIP|以上

## 重要
- 必ず入力と同じ段落数を出力すること（空行も含めて1行ずつ対応）
- 空段落は「SKIP|」と出力
- テキスト内容は変更せず、見出し番号の付与・修正のみ行う
- 全角数字を使用する"""


# --tone 付きプロンプト: 構造化＋公用文変換を一括実行
AUTO_TONE_SYSTEM_PROMPT = """\
あなたは日本の裁判所提出書面の書式整形と公用文変換を行う専門家です。
与えられた文章の各段落を分析し、(1) 裁判所書式の見出し階層を判定し、(2) 本文を公用文に変換してください。

## タスク1: 見出し階層判定

- H1: 最上位の大見出し → 「第１　」「第２　」… を付与
- H2: 中見出し → 「１　」「２　」… を付与
- H3: 小見出し → 「(1)　」「(2)　」… を付与
- H4: 細目 → 「ア　」「イ　」… を付与
- H5〜H7: 同様に階層化
- BODY: 本文（見出しではない）
- SKIP: 冒頭部分（事件番号・当事者名・表題・日付・署名等）・末尾の「以上」等

### 構造判定ルール
1. 元の文章に見出し番号が既にあれば、正しい階層の全角番号に置き換える
2. 意味的に見出しと判断できれば番号を振る
3. 各階層の番号は出現順に１から連番（H2はH1ごとにリセット、H3はH2ごとにリセット…）

## タスク2: 公用文変換（BODY段落のみに適用）

文化審議会「公用文作成の考え方」および本多勝一「日本語の作文技術」に準拠して、BODY段落のテキストを変換する。

### 共通ルール（全tone）

- 句読点: 「。」「、」に統一（カンマは読点に）
- 冗長表現の排除: 「利用することができる→利用できる」「調査を実施した→調査した」「についてであるが→について」
- 二重否定の排除: 「しないわけではない→することもある」
- 不必要な受身形を避ける
- 一文を短くする（一文一論点）
- 修飾語の順序: 長い修飾語を先に、短い修飾語を後に
- 係る語と受ける語を近づける
- 同じ助詞（特に「の」「が」「を」）の連続使用を避ける
- 文末表現の単調さを避ける（「〜である」の連続等）

### tone別ルール
{tone_rules}

### 変換時の禁止事項
- 法令の引用・号証番号は変更しない
- 固有名詞は変更しない
- 日付・数値の事実は変更しない
- 意味を変えない範囲で表現を整える

## 出力形式
各段落を1行ずつ:
TYPE|テキスト

- SKIP段落: 元のテキストをそのまま出力
- 見出し段落: 番号を付与したテキストを出力（見出しテキスト自体は公用文変換しない）
- BODY段落: 公用文変換済みのテキストを出力

## 重要
- 必ず入力と同じ段落数を出力すること（空行も含めて1行ずつ対応）
- 空段落は「SKIP|」と出力
- 全角数字を使用する"""


TONE_RULES = {
    'strict': """\
**strict（厳格）** — 法令・告示準拠:
- 常体（である体）を厳守
- 法令特有の用語を正確に使用（及び/並びに/又は/若しくは の使い分け）
- 「べく」「べし」は使わず「べきである」
- 文語の名残を排除（〜のごとく→〜のように）
- 修飾は最小限、事実と法的評価のみ""",

    'formal': """\
**formal（フォーマル）** — 準備書面・答弁書向け:
- 常体（である体）
- 法律用語を適切に使用しつつ、過度な硬さは避ける
- 論理的接続を明確に（したがって、もっとも、これに対し、なお）
- 引用・参照を正確に（「同条同項」「前記第１の２」等）
- 事実→評価→結論の論理展開を明確にする""",

    'polite': """\
**polite（丁寧）** — 陳述書・上申書向け:
- 敬体（です・ます体）で統一
- 「ございます」は使わない
- 「おります」「いたします」は必要最小限
- 主語を明確にし、経験・心情を具体的に記述""",

    'plain': """\
**plain（平易）** — 依頼者向け説明:
- 敬体（です・ます体）
- 専門用語は必ず説明を付けるか言い換える
- 義務教育レベルの知識で理解できるように
- 短い文、具体的な例示を多用""",
}


# 連番ミス修正用プロンプト
FIX_NUMBERING_SYSTEM_PROMPT = """\
あなたは日本の裁判所提出書面の見出し番号の校正を行う専門家です。
与えられた文章の各段落を分析し、見出し番号の誤り（連番飛び・レベルまたぎ・書式混在）を検出して修正してください。

## 岡口マクロ準拠の正しい番号体系

| レベル | 番号形式 | 例 |
|---|---|---|
| Level 1 | 第＋全角数字 | 第１、第２、第３ |
| Level 2 | 全角数字 | １、２、３ |
| Level 3 | 半角括弧＋半角数字 | (1)、(2)、(3) |
| Level 4 | 全角カタカナ | ア、イ、ウ |
| Level 5 | 半角括弧＋半角カタカナ | (ｱ)、(ｲ)、(ｳ) |
| Level 6 | 全角小文字英字 | ａ、ｂ、ｃ |
| Level 7 | 半角括弧＋半角小文字英字 | (a)、(b)、(c) |

## 検出・修正すべきパターン

1. **連番飛び**: 「１」→「３」（２が抜けている）→ 文脈判断で修正
2. **レベルまたぎ**: 「第１」の下に突然「(1)」（Level 2の「１」が抜けている）→ 文脈から正しいレベルを判定
3. **書式混在**: 同一レベルで「１」「⑵」「(3)」が混在 → 正しい書式に統一
4. **丸数字の修正**: ⑴⑵⑶ → (1)(2)(3)、①②③ → 文脈からレベル判定して正しい番号に
5. **全角半角混在**: （１）と(1)が混在 → Level 3は半角括弧＋半角数字に統一
6. **番号リセット忘れ**: H1が変わったのにH2の番号がリセットされていない

## ルール
- 番号のない段落（BODY/SKIP）はそのまま返す
- 見出しテキスト本体は変更しない（番号部分のみ修正）
- 修正があった場合、修正内容をコメントとして記載

## 出力形式
各段落を1行ずつ:
TYPE|テキスト|COMMENT（修正があった場合のみCOMMENTを追加）

例:
SKIP|令和６年（ワ）第１２３号
H1|第１　被告の主張|
H2|１　契約成立|
BODY|原告と被告は…|
H2|２　債務不履行|FIXED: ３→２（連番飛び修正）
H3|(1)　履行遅滞|FIXED: ⑴→(1)（丸括弧→半角括弧）
BODY|被告は…|
H3|(2)　損害|FIXED: ⑶→(2)（連番飛び＋書式修正）

## 重要
- 必ず入力と同じ段落数を出力すること
- 空段落は「SKIP|」と出力
- 修正理由は必ずCOMMENTに記載
- 修正がない段落はCOMMENT部分を省略してよい
- 全角数字を使用する"""

# H1〜H7のカタカナ順序
_KATAKANA_ORDER = 'アイウエオカキクケコサシスセソタチツテトナニヌネノハヒフヘホマミムメモヤユヨラリルレロワ'
_HAN_KATA_ORDER = 'ｱｲｳｴｵｶｷｸｹｺｻｼｽｾｿﾀﾁﾂﾃﾄﾅﾆﾇﾈﾉﾊﾋﾌﾍﾎﾏﾐﾑﾒﾓﾔﾕﾖﾗﾘﾙﾚﾛﾜ'
_ALPHA_ORDER = 'ａｂｃｄｅｆｇｈｉｊｋｌｍｎｏｐｑｒｓｔｕｖｗｘｙｚ'
_HAN_ALPHA_ORDER = 'abcdefghijklmnopqrstuvwxyz'

# 全角数字変換
_ZENKAKU_DIGITS = '０１２３４５６７８９'
def _to_zenkaku_num(n):
    """整数を全角数字文字列に変換。"""
    return ''.join(_ZENKAKU_DIGITS[int(d)] for d in str(n))


# ============================================================
# 事実摘示チェッカー（--check モード）
# ============================================================

CHECK_SYSTEM_PROMPT = """\
あなたは日本の裁判所提出書面の校正チェッカーです。
与えられた文章の各段落を分析し、事実摘示の構文ルールに違反している箇所を指摘してください。

## チェック項目

### 1. 主語の欠落（最重要）
法律文書では全ての文に主語が必要である。主語（「〜は」「〜が」）のない文を検出する。

検出パターン:
- 文頭がいきなり述語や目的語で始まる（例: 「支払った。」「契約を締結した。」）
- 接続詞の後に主語なく述語が続く（例: 「しかし、支払わなかった。」→ 誰が？）
- 受身形で行為主体が不明（例: 「解除された。」→ 誰に？）
- 連文節の途中で主語がすり替わっている

除外:
- 見出し行は主語不要
- 「以上」「記」等の定型句
- 「なお」「ただし」で始まる補足文で、直前の文の主語が明らかに引き継がれる場合

### 2. 構文順序の逸脱
事実を記載する文の標準的な語順は:
  主体（は）→ 相手方（に対し）→ 日時 → 目的物 → 法律行為

以下のパターンを検出する:
- 日時が文末近くにある（例: 「原告は契約を締結した、令和５年１月に。」）
- 相手方が目的物の後にある（例: 「原告は土地を被告に売った。」→「原告は、被告に対し、土地を売った。」）
- 目的物の説明が法律行為の後にある

ただし:
- 強調のための語順変更は指摘しつつ許容とする
- 評価や法的主張の文（「〜と解すべきである」等）は事実摘示ではないので語順チェック対象外

## 出力形式
各段落を1行ずつ、以下の形式で出力:
行番号|指摘内容

指摘がない段落は出力しない。
1つの段落に複数の指摘がある場合は複数行出力する。

指摘内容のフォーマット:
- [主語欠落] 「〜した。」→ 主語がありません。誰が〜したのか明記してください。
- [語順] 「原告は土地を被告に売った。」→ 「原告は、被告に対し、土地を売った。」（相手方を目的物の前に）
- [主語交替] 「原告は契約を締結し、支払った。」→ 2つ目の述語「支払った」の主語は原告で正しいですか？

## 重要
- 指摘は具体的に。該当する文（またはその一部）を引用すること
- 修正案がある場合は「→」の後に示す
- 見出し行・SKIP行は対象外
- 過度な指摘は避ける。明らかな問題のみ指摘する"""


def check_writing_with_llm(paragraphs_text, provider=None, model=None):
    """LLMで事実摘示の構文チェックを行う。"""
    n = len(paragraphs_text)
    user_content = (
        f"以下の{n}段落の裁判所提出書面をチェックしてください。\n"
        f"主語の欠落と構文順序の逸脱を指摘してください。\n\n"
    )
    for i, text in enumerate(paragraphs_text):
        if text.strip():
            user_content += f"段落{i+1}: {text.strip()}\n"

    result_text = _call_llm(CHECK_SYSTEM_PROMPT, user_content,
                            provider=provider, model=model)

    # パース: 行番号|指摘内容
    findings = []
    for line in result_text.strip().split('\n'):
        line = line.strip()
        if not line or '|' not in line:
            continue
        parts = line.split('|', 1)
        try:
            # 「段落3」「3」等の形式に対応
            num_str = parts[0].strip().replace('段落', '').replace('行', '')
            para_num = int(num_str)
            finding = parts[1].strip()
            findings.append((para_num, finding))
        except ValueError:
            continue

    return findings


def run_check(input_path, provider=None, model=None, add_comments=True, output_path=None):
    """事実摘示チェッカーを実行。指摘をWordコメントとして追加。"""
    from docx import Document as Doc

    doc = Doc(input_path)
    para_texts = [p.text for p in doc.paragraphs]

    provider_name = provider or os.environ.get('COURT_FORMAT_LLM_PROVIDER', 'anthropic')
    print(f"段落数: {len(para_texts)}")
    print(f"事実摘示チェック中... (provider: {provider_name})")

    findings = check_writing_with_llm(para_texts, provider=provider, model=model)

    if not findings:
        print("\n指摘事項はありません。")
        return input_path, []

    print(f"\n指摘事項: {len(findings)} 件")
    for para_num, finding in findings:
        print(f"  段落{para_num}: {finding}")

    # Wordコメントとして書き込み
    if add_comments and output_path:
        # 段落番号→docインデックスの対応（非空段落の番号）
        nonempty_map = {}
        count = 0
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip():
                count += 1
                nonempty_map[count] = i

        comment_count = 0
        for para_num, finding in findings:
            doc_idx = nonempty_map.get(para_num)
            if doc_idx is not None and doc_idx < len(doc.paragraphs):
                add_comment_to_paragraph(
                    doc, doc.paragraphs[doc_idx],
                    f"【校正】{finding}",
                    author="court-format-checker"
                )
                comment_count += 1

        if output_path is None:
            base, ext = os.path.splitext(input_path)
            output_path = f"{base}_チェック済{ext}"

        doc.save(output_path)
        print(f"\nコメント: {comment_count}箇所に指摘を記録")
        print(f"出力: {output_path}")
        return output_path, findings

    return input_path, findings


def _load_api_key():
    """ANTHROPIC_API_KEY を環境変数 or .env ファイルから取得。"""
    key = os.environ.get('ANTHROPIC_API_KEY')
    if key:
        return key
    # .env ファイルから探す
    env_paths = [
        os.path.expanduser('~/dev/legal-jp-local/.env'),
        os.path.expanduser('~/dev/legal-jp-mcp-server/.env'),
        os.path.expanduser('~/.env'),
    ]
    for path in env_paths:
        if os.path.exists(path):
            with open(path) as f:
                for line in f:
                    if line.startswith('ANTHROPIC_API_KEY='):
                        return line.strip().split('=', 1)[1]
    return None


def _call_llm(system_prompt, user_content, provider=None, model=None):
    """LLM呼び出しの共通関数。Anthropic API / Ollama / OpenAI互換に対応。

    provider: 'anthropic' (default), 'ollama', 'openai'
    model: モデル名（省略時はproviderに応じたデフォルト）
    """
    if provider is None:
        # 環境変数で切り替え可能
        provider = os.environ.get('COURT_FORMAT_LLM_PROVIDER', 'anthropic')

    if provider == 'ollama':
        return _call_ollama(system_prompt, user_content, model)
    elif provider == 'openai':
        return _call_openai_compat(system_prompt, user_content, model)
    else:
        return _call_anthropic(system_prompt, user_content, model)


def _call_anthropic(system_prompt, user_content, model=None):
    """Anthropic API 呼び出し。"""
    import anthropic

    api_key = _load_api_key()
    if not api_key:
        print("エラー: ANTHROPIC_API_KEY が見つかりません。")
        print("環境変数に設定するか、~/dev/legal-jp-local/.env に記載してください。")
        sys.exit(1)

    if model is None:
        model = os.environ.get('COURT_FORMAT_MODEL', 'claude-haiku-4-5-20251001')

    client = anthropic.Anthropic(api_key=api_key)
    message = client.messages.create(
        model=model,
        max_tokens=8192,
        system=system_prompt,
        messages=[{"role": "user", "content": user_content}],
    )
    return message.content[0].text


def _call_ollama(system_prompt, user_content, model=None):
    """Ollama (ローカルLLM) 呼び出し。"""
    import json
    import urllib.request

    if model is None:
        model = os.environ.get('COURT_FORMAT_MODEL', 'qwen2.5:14b')

    base_url = os.environ.get('OLLAMA_HOST', 'http://localhost:11434')
    url = f"{base_url}/api/chat"

    payload = json.dumps({
        "model": model,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_content},
        ],
        "stream": False,
        "options": {"num_predict": 8192},
    }).encode('utf-8')

    req = urllib.request.Request(url, data=payload,
                                headers={"Content-Type": "application/json"})
    try:
        with urllib.request.urlopen(req, timeout=300) as resp:
            result = json.loads(resp.read().decode('utf-8'))
            return result['message']['content']
    except Exception as e:
        print(f"エラー: Ollama接続失敗: {e}")
        print(f"Ollamaが起動しているか確認してください: {base_url}")
        sys.exit(1)


def _call_openai_compat(system_prompt, user_content, model=None):
    """OpenAI互換API呼び出し（LM Studio, vLLM等）。"""
    import json
    import urllib.request

    if model is None:
        model = os.environ.get('COURT_FORMAT_MODEL', 'default')

    base_url = os.environ.get('OPENAI_BASE_URL', 'http://localhost:1234/v1')
    url = f"{base_url}/chat/completions"

    payload = json.dumps({
        "model": model,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_content},
        ],
        "max_tokens": 8192,
    }).encode('utf-8')

    api_key = os.environ.get('OPENAI_API_KEY', 'no-key')
    req = urllib.request.Request(url, data=payload, headers={
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}",
    })
    try:
        with urllib.request.urlopen(req, timeout=300) as resp:
            result = json.loads(resp.read().decode('utf-8'))
            return result['choices'][0]['message']['content']
    except Exception as e:
        print(f"エラー: OpenAI互換API接続失敗: {e}")
        print(f"サーバーが起動しているか確認してください: {base_url}")
        sys.exit(1)


def _parse_llm_output(result_text, expect_comments=False):
    """LLM出力をパースして (tag, text) または (tag, text, comment) のリストを返す。"""
    lines = result_text.strip().split('\n')
    parsed = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if '|' in line:
            parts = line.split('|')
            tag = parts[0].strip().upper()
            text = parts[1] if len(parts) > 1 else ''
            if expect_comments:
                comment = parts[2].strip() if len(parts) > 2 else ''
                parsed.append((tag, text, comment))
            else:
                parsed.append((tag, text))
    return parsed


def analyze_with_llm(paragraphs_text, tone=None, provider=None, model=None):
    """LLMで段落構造を解析する。tone指定時は公用文変換も同時実行。"""
    n = len(paragraphs_text)

    if tone and tone in TONE_RULES:
        system = AUTO_TONE_SYSTEM_PROMPT.format(tone_rules=TONE_RULES[tone])
        user_content = (
            f"以下の{n}段落を分析し、見出し階層を判定し、"
            f"BODY段落は「{tone}」toneで公用文変換してください。\n"
            f"必ず{n}行出力してください。\n\n"
        )
    else:
        system = AUTO_SYSTEM_PROMPT
        user_content = (
            f"以下の{n}段落を分析し、見出し階層を判定してください。\n"
            f"必ず{n}行出力してください。\n\n"
        )

    for i, text in enumerate(paragraphs_text):
        user_content += f"段落{i+1}: {text}\n"

    result_text = _call_llm(system, user_content, provider=provider, model=model)
    parsed = _parse_llm_output(result_text)

    # 段落数ミスマッチの警告
    if len(parsed) != n:
        print(f"警告: 入力{n}段落 vs LLM出力{len(parsed)}行（ミスマッチ）")

    return parsed


def fix_numbering_with_llm(paragraphs_text, provider=None, model=None):
    """LLMで連番ミスを検出・修正する。修正コメント付きで返す。"""
    user_content = (
        "以下の文章の各段落を分析し、見出し番号の誤りを検出・修正してください。\n\n"
    )
    for i, text in enumerate(paragraphs_text):
        user_content += f"段落{i+1}: {text}\n"

    result_text = _call_llm(FIX_NUMBERING_SYSTEM_PROMPT, user_content,
                            provider=provider, model=model)
    return _parse_llm_output(result_text, expect_comments=True)


def _apply_structure_to_doc(doc, structure, add_comments=False):
    """構造解析結果（タグ+テキストのリスト）をdocに適用する共通処理。
    見出し番号・インデント・フォント・ページ設定を適用する。
    BODY段落のテキストは書き換えない。add_comments=True の場合、
    LLMが修正候補を出した箇所にWordコメントで修正案を残す。"""
    setup_page(doc)
    setup_default_style(doc)

    level_map = {'H1': 1, 'H2': 2, 'H3': 3, 'H4': 4, 'H5': 5, 'H6': 6, 'H7': 7}
    current_heading_level = 0
    comment_count = 0

    # structureから空SKIP行を除いて、doc非空段落と1:1対応させるマップを構築
    struct_nonempty = [(tag, text) for tag, text in structure
                       if not (tag == 'SKIP' and not text.strip())]
    doc_nonempty_indices = [i for i, p in enumerate(doc.paragraphs) if p.text.strip()]
    para_map = {}  # {doc段落index: (tag, new_text)}
    for idx, doc_i in enumerate(doc_nonempty_indices):
        if idx < len(struct_nonempty):
            para_map[doc_i] = struct_nonempty[idx]

    for para_idx, para in enumerate(doc.paragraphs):
        # 半角→全角変換
        for run in para.runs:
            convert_run_to_zenkaku(run)
        # フォント統一
        set_paragraph_font(para, size=12)

        text = para.text.strip()
        if not text:
            continue

        # マッピングから構造情報を取得
        if para_idx in para_map:
            tag, new_text = para_map[para_idx]
        else:
            tag, new_text = 'BODY', text

        if tag == 'SKIP':
            if SKIP_PATTERNS.match(text):
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                set_indent(para)
            continue
        elif tag in level_map:
            level = level_map[tag]
            current_heading_level = level
            # 見出し番号は書き換える（構造化の本質）
            if new_text.strip():
                for i, run in enumerate(para.runs):
                    if i == 0:
                        run.text = new_text
                    else:
                        run.text = ''
            set_heading_indent(para, level)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            # BODY段落: 原文は書き換えない。修正候補があればコメントで残す
            new_stripped = new_text.strip()
            if add_comments and new_stripped and new_stripped != text:
                add_comment_to_paragraph(
                    doc, para,
                    f"【修正候補】{new_stripped}",
                    author="court-format"
                )
                comment_count += 1

            set_body_indent(para, current_heading_level)
            if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    format_tables(doc)
    add_page_number(doc)
    return comment_count


def convert_auto(input_path, output_path=None, tone=None, provider=None, model=None,
                 show_comments=None):
    """LLMで構造解析してから裁判所書式に変換（--auto モード）。
    tone指定時は公用文変換も同時実行。
    show_comments: True=変更箇所にWordコメントで元テキストを残す。
                   None（デフォルト）=tone指定時は自動でTrue。"""
    if output_path is None:
        base, ext = os.path.splitext(input_path)
        output_path = f"{base}_裁判所書式{ext}"

    # tone指定時はデフォルトでコメントON
    if show_comments is None:
        show_comments = tone is not None

    doc = Document(input_path)
    para_texts = [p.text for p in doc.paragraphs]

    provider_name = provider or os.environ.get('COURT_FORMAT_LLM_PROVIDER', 'anthropic')
    print(f"段落数: {len(para_texts)}")
    print(f"LLM構造解析中... (provider: {provider_name})")
    if tone:
        print(f"公用文変換: tone={tone}")
    if show_comments:
        print("コメント: ON（修正候補をWordコメントで残します。原文は変更しません）")

    analysis = analyze_with_llm(para_texts, tone=tone, provider=provider, model=model)
    print(f"LLM解析結果: {len(analysis)} 段落")

    comment_count = _apply_structure_to_doc(doc, analysis, add_comments=show_comments)
    doc.save(output_path)

    if show_comments and comment_count > 0:
        print(f"コメント: {comment_count}箇所に修正候補を記録")

    return output_path


def convert_fix_numbering(input_path, output_path=None, provider=None, model=None,
                          show_comments=True):
    """LLMで連番ミスを検出・修正してから裁判所書式に変換（--fix-numbering モード）。
    show_comments: デフォルトTrue。修正箇所にWordコメントで修正理由を残す。"""
    if output_path is None:
        base, ext = os.path.splitext(input_path)
        output_path = f"{base}_裁判所書式{ext}"

    doc = Document(input_path)
    para_texts = [p.text for p in doc.paragraphs]

    provider_name = provider or os.environ.get('COURT_FORMAT_LLM_PROVIDER', 'anthropic')
    print(f"段落数: {len(para_texts)}")
    print(f"連番ミス検出中... (provider: {provider_name})")

    analysis = fix_numbering_with_llm(para_texts, provider=provider, model=model)

    # 修正コメントを表示
    fixes = [(tag, text, comment) for tag, text, comment in analysis if comment]
    if fixes:
        print(f"\n修正箇所: {len(fixes)} 件")
        for tag, text, comment in fixes:
            print(f"  {comment}")
            display_text = text[:40] + '...' if len(text) > 40 else text
            print(f"    → {tag}|{display_text}")
    else:
        print("\n連番ミスは検出されませんでした。")

    # (tag, text, comment) → (tag, text) に変換して適用
    structure = [(tag, text) for tag, text, comment in analysis]
    _apply_structure_to_doc(doc, structure, add_comments=False)

    # fix-numberingの場合はコメントで修正理由を残す（構造適用とは別）
    if show_comments and fixes:
        # 修正箇所にコメント追加
        fix_comments = {text: comment for tag, text, comment in fixes}
        for para in doc.paragraphs:
            text = para.text.strip()
            if text in fix_comments:
                add_comment_to_paragraph(
                    doc, para,
                    f"【連番修正】{fix_comments[text]}",
                    author="court-format"
                )

    doc.save(output_path)
    return output_path, fixes


def convert_with_structure(input_path, structure_json_path, output_path=None):
    """JSONファイルの構造定義を使って裁判所書式に変換（--apply-structure モード）。
    Claude Codeのスキルから呼ばれる。API不要。"""
    import json

    if output_path is None:
        base, ext = os.path.splitext(input_path)
        output_path = f"{base}_裁判所書式{ext}"

    with open(structure_json_path, 'r') as f:
        structure_data = json.load(f)

    # JSON形式: [{"index": 0, "tag": "H1", "text": "第１　..."}, ...]
    structure = [(item['tag'].upper(), item.get('text', '')) for item in structure_data]

    doc = Document(input_path)
    _apply_structure_to_doc(doc, structure)
    doc.save(output_path)
    return output_path


# ============================================================
# メイン変換処理
# ============================================================

def _extract_paragraphs_and_tables(input_path):
    """元のdocxからテキストとテーブルデータを抽出する。
    Returns: list of dict {
        type: 'para'|'table',
        text: str,
        alignment: str|None,
        fmt: dict|None,
        rows: list|None,
    }
    段落とテーブルの出現順序を保持する。"""
    src = Document(input_path)
    elements = []

    # docx内部ではw:bodyの直下にw:pとw:tblが混在している
    body = src.element.body
    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            # 段落
            text = ''.join(node.text or '' for node in child.iter(qn('w:t')))
            # 配置を取得
            pPr = child.find(qn('w:pPr'))
            alignment = None
            spacing_before = None
            spacing_after = None
            line = None
            line_rule = None
            if pPr is not None:
                jc = pPr.find(qn('w:jc'))
                if jc is not None:
                    alignment = jc.get(qn('w:val'))
                spacing = pPr.find(qn('w:spacing'))
                if spacing is not None:
                    spacing_before = spacing.get(qn('w:before'))
                    spacing_after = spacing.get(qn('w:after'))
                    line = spacing.get(qn('w:line'))
                    line_rule = spacing.get(qn('w:lineRule'))
            elements.append({
                'type': 'para',
                'text': text,
                'alignment': alignment,
                'fmt': {
                    'space_before': spacing_before,
                    'space_after': spacing_after,
                    'line': line,
                    'line_rule': line_rule,
                    'bold': _is_fully_bold(child),
                },
            })
        elif tag == 'tbl':
            # テーブル: セルのテキストを抽出
            rows = []
            for tr in child.findall(qn('w:tr')):
                cells = []
                for tc in tr.findall(qn('w:tc')):
                    cell_text = ''.join(
                        node.text or '' for node in tc.iter(qn('w:t'))
                    )
                    cells.append(cell_text)
                rows.append(cells)
            elements.append({'type': 'table', 'rows': rows})

    return elements


# 岡口マクロが設定するスタイル名。これらが検出されたらインデント処理をスキップ。
OKAGUCHI_STYLES = re.compile(r'^(ランク[１-９1-9]|本文[１-９1-9]|標準\(太郎文書スタイル\))')


def _has_okaguchi_styles(input_path):
    """元の文書に岡口マクロのスタイルが使われているか判定。"""
    doc = Document(input_path)
    for para in doc.paragraphs:
        if para.style and OKAGUCHI_STYLES.match(para.style.name):
            return True
    return False


def convert(input_path, output_path=None):
    """docxファイルを裁判所書式に変換。
    新規docxを作成してテキストを流し込む方式。元のスタイル・グリッドに依存しない。
    岡口マクロのスタイルが検出されたらインデント処理をスキップ。"""

    if output_path is None:
        base, ext = os.path.splitext(input_path)
        output_path = f"{base}_裁判所書式{ext}"

    # 0. 岡口マクロ検出
    skip_indent = _has_okaguchi_styles(input_path)
    if skip_indent:
        print("注意: 岡口マクロのスタイルが検出されました。インデント処理をスキップします。")

    # 1. 元のdocxからテキスト・テーブルを抽出
    elements = _extract_paragraphs_and_tables(input_path)

    # 2. 新規docxを作成
    doc = Document()

    # 3. ページ設定・デフォルトスタイル
    setup_page(doc)
    setup_default_style(doc)

    # 4. 最初の空段落を削除（Documentが自動生成する）
    if doc.paragraphs:
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

    # 5. テキストを流し込み
    current_heading_level = 0
    in_header_section = True

    for elem in elements:
        if elem['type'] == 'table':
            # テーブルを追加
            rows = elem['rows']
            if not rows:
                continue
            max_cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=max_cols)
            for ri, row_data in enumerate(rows):
                for ci, cell_text in enumerate(row_data):
                    if ci < max_cols:
                        cell = table.cell(ri, ci)
                        cell.text = to_zenkaku(cell_text)
                        for p in cell.paragraphs:
                            set_paragraph_font(p, size=12)
            format_tables(doc)
            continue

        # 段落
        raw_text = elem['text']
        text = to_zenkaku(raw_text).strip()

        # 空段落
        if not text:
            para = doc.add_paragraph('')
            apply_paragraph_layout(para, elem.get('fmt'))
            if elem.get('alignment') == 'right':
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif elem.get('alignment') == 'center':
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        # 段落を追加
        para = doc.add_paragraph(text)
        set_paragraph_font(para, size=12)
        apply_paragraph_layout(para, elem.get('fmt'))
        if elem.get('alignment') == 'right':
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif elem.get('alignment') == 'center':
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 冒頭セクション判定
        if in_header_section:
            level = detect_heading_level(text)
            if level is not None:
                in_header_section = False
            elif is_header_section(text):
                # 冒頭セクション: 元の配置を維持
                if elem.get('alignment') == 'right':
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif elem.get('alignment') == 'center':
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue
            else:
                continue

        # 見出し・本文判定
        if not in_header_section:
            level = detect_heading_level(text)

            if level is not None:
                current_heading_level = level
                # 番号後のスペースを全角1個に正規化
                normalized = normalize_heading_spacing(text)
                if normalized != text:
                    para.runs[0].text = normalized
                if not skip_indent:
                    set_heading_indent(para, level)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                if SKIP_PATTERNS.match(text):
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    if not skip_indent:
                        set_indent(para)
                else:
                    if not skip_indent:
                        set_body_indent(para, current_heading_level)
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 6. ページ番号
    add_page_number(doc)

    # 7. 保存
    doc.save(output_path)
    return output_path


# ============================================================
# CLI
# ============================================================

if __name__ == '__main__':
    import argparse

    parser = argparse.ArgumentParser(
        description='裁判所書式変換ツール（岡口マクロ準拠）',
        epilog="""
LLMプロバイダーの切り替え:
  環境変数 COURT_FORMAT_LLM_PROVIDER で指定（anthropic / ollama / openai）
  または --provider オプションで直接指定

  例:
    # Anthropic API（デフォルト）
    python3 court_format_converter.py --auto --tone formal input.docx

    # Ollama（ローカルLLM）
    python3 court_format_converter.py --auto --provider ollama --model qwen2.5:14b input.docx

    # OpenAI互換API（LM Studio等）
    python3 court_format_converter.py --auto --provider openai input.docx
""",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument('input', help='入力docxファイル')
    parser.add_argument('output', nargs='?', default=None,
                        help='出力docxファイル（省略時は_裁判所書式.docxを付加）')
    parser.add_argument('--auto', action='store_true',
                        help='LLMで構造解析し、見出し番号を自動付与')
    parser.add_argument('--tone', choices=['strict', 'formal', 'polite', 'plain'],
                        help='公用文変換のtone（--autoと併用）。'
                             'strict=法令準拠, formal=準備書面, polite=陳述書, plain=平易')
    parser.add_argument('--fix-numbering', action='store_true',
                        help='LLMで連番ミス（飛び・レベルまたぎ・書式混在）を検出・修正')
    parser.add_argument('--check', action='store_true',
                        help='事実摘示チェッカー: 主語欠落・構文順序をLLMで検出しWordコメントで指摘')
    parser.add_argument('--apply-structure', metavar='JSON',
                        help='Claude Codeが生成した構造JSONファイルを適用（API不要）')
    parser.add_argument('--provider', choices=['anthropic', 'ollama', 'openai'],
                        help='LLMプロバイダー（デフォルト: anthropic）')
    parser.add_argument('--model',
                        help='LLMモデル名（省略時はプロバイダーのデフォルト）')
    parser.add_argument('--no-comments', action='store_true',
                        help='公用文変換・連番修正時にWordコメントを付けない')

    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"エラー: ファイルが見つかりません: {args.input}")
        sys.exit(1)

    if args.tone and not args.auto and not args.check:
        print("警告: --tone は --auto と併用してください。--auto を有効にします。")
        args.auto = True

    if args.check:
        out = args.output
        if out is None:
            base, ext = os.path.splitext(args.input)
            out = f"{base}_チェック済{ext}"
        result, findings = run_check(
            args.input,
            provider=args.provider, model=args.model,
            add_comments=not args.no_comments,
            output_path=out)
        if findings:
            print(f"\n合計: {len(findings)} 件の指摘")
        sys.exit(0)
    elif args.apply_structure:
        if not os.path.exists(args.apply_structure):
            print(f"エラー: JSONファイルが見つかりません: {args.apply_structure}")
            sys.exit(1)
        result = convert_with_structure(args.input, args.apply_structure, args.output)
        print(f"変換完了: {result}")
    elif args.fix_numbering:
        result, fixes = convert_fix_numbering(
            args.input, args.output,
            provider=args.provider, model=args.model,
            show_comments=not args.no_comments)
        print(f"\n変換完了: {result}")
        if fixes:
            print(f"修正件数: {len(fixes)}")
    elif args.auto:
        show_comments = None if not args.no_comments else False
        result = convert_auto(
            args.input, args.output, tone=args.tone,
            provider=args.provider, model=args.model,
            show_comments=show_comments)
        print(f"変換完了: {result}")
    else:
        result = convert(args.input, args.output)
        print(f"変換完了: {result}")
